import asyncio
import aiofiles
import tempfile
import os
from typing import List, Dict, Any, Optional
from urllib.parse import urlparse
import structlog
import httpx
from io import BytesIO
import re # Import re at the top level

import PyPDF2
import pdfplumber
from docx import Document
import email
from email.mime.text import MIMEText

# Assuming settings and schemas are defined elsewhere correctly
from app.core.config import settings
from app.models.schemas import DocumentMetadata, EmbeddingChunk

# Mock objects for stand-alone execution
class MockSettings:
    chunk_size = 512
logger = structlog.get_logger(__name__)

class DocumentProcessor:
    """Handles document ingestion and processing for various file formats."""

    def __init__(self):
        self.chunk_size = settings.chunk_size
        self.chunk_overlap = settings.chunk_overlap
        
        # Cost-optimized chunking configuration
        self.use_semantic_chunking = True
        self.max_paragraph_chunk_size = getattr(settings, 'max_paragraph_chunk_size', self.chunk_size * 1.25)
        self.min_chunk_size = getattr(settings, 'min_chunk_size', 100)
        self.preserve_section_boundaries = True
        
        # Cost control settings
        self.max_tokens_per_chunk = 800
        self.enable_chunk_compression = True

    async def download_document(self, url: str) -> bytes:
        """Download document from URL."""
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                response = await client.get(url)
                response.raise_for_status()
                return response.content
        except Exception as e:
            logger.error("Failed to download document", url=url, error=str(e))
            raise ValueError(f"Failed to download document: {str(e)}")

    async def process_document(self, url: str) -> tuple[DocumentMetadata, List[EmbeddingChunk]]:
        """Process document from URL and return metadata and chunks."""
        logger.info("Processing document", url=url)
        
        # Download document
        document_content = await self.download_document(url)
        
        # Determine file type from URL
        parsed_url = urlparse(url)
        file_extension = os.path.splitext(parsed_url.path)[1].lower()
        
        if not file_extension:
            # Try to determine from content type or content
            if document_content.startswith(b'%PDF'):
                file_extension = '.pdf'
            elif document_content.startswith(b'PK'):
                file_extension = '.docx'
            else:
                file_extension = '.txt'
        
        with tempfile.NamedTemporaryFile(suffix=file_extension, delete=False) as temp_file:
            temp_file.write(document_content)
            temp_file_path = temp_file.name
        
        try:
            if file_extension == '.pdf':
                text_content, total_pages = await self._process_pdf(temp_file_path)
            elif file_extension == '.docx':
                text_content, total_pages = await self._process_docx(temp_file_path)
            elif file_extension in ['.eml', '.msg']:
                text_content, total_pages = await self._process_email(temp_file_path)
            else:
                async with aiofiles.open(temp_file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text_content = await f.read()
                total_pages = None
            
            chunks = self._create_chunks(text_content, url)
            
            document_id = self._generate_document_id(url)
            metadata = DocumentMetadata(
                document_id=document_id,
                document_type=file_extension[1:],
                total_pages=total_pages,
                total_chunks=len(chunks),
                processing_time=0.0
            )
            
            logger.info("Document processed successfully", 
                        document_id=document_id, 
                        chunks=len(chunks), 
                        pages=total_pages)
            
            return metadata, chunks
            
        finally:
            os.unlink(temp_file_path)

    async def _process_pdf(self, file_path: str) -> tuple[str, int]:
        """Process PDF file and extract text."""
        text_content = ""
        total_pages = 0
        
        try:
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
        except Exception as e:
            logger.warning("pdfplumber failed, trying PyPDF2", error=str(e))
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    total_pages = len(pdf_reader.pages)
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            except Exception as e2:
                logger.error("Both PDF processors failed", error=str(e2))
                raise ValueError(f"Failed to process PDF: {str(e2)}")
        
        return text_content.strip(), total_pages
    
    async def _process_docx(self, file_path: str) -> tuple[str, Optional[int]]:
        """Process DOCX file and extract text."""
        try:
            doc = Document(file_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_content += row_text + "\n"
            
            return text_content.strip(), None
            
        except Exception as e:
            logger.error("Failed to process DOCX", error=str(e))
            raise ValueError(f"Failed to process DOCX: {str(e)}")
    
    async def _process_email(self, file_path: str) -> tuple[str, Optional[int]]:
        """Process email file and extract text."""
        try:
            with open(file_path, 'rb') as f:
                msg = email.message_from_bytes(f.read())
            
            text_content = ""
            
            text_content += f"From: {msg.get('From', 'Unknown')}\n"
            text_content += f"To: {msg.get('To', 'Unknown')}\n"
            text_content += f"Subject: {msg.get('Subject', 'No Subject')}\n"
            text_content += f"Date: {msg.get('Date', 'Unknown')}\n\n"
            
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        text_content += part.get_payload(decode=True).decode('utf-8', errors='ignore')
            else:
                text_content += msg.get_payload(decode=True).decode('utf-8', errors='ignore')
            
            return text_content.strip(), None
            
        except Exception as e:
            logger.error("Failed to process email", error=str(e))
            raise ValueError(f"Failed to process email: {str(e)}")

    def _create_chunks(self, text: str, document_url: str) -> List[EmbeddingChunk]:
        """Create chunks from text content using cost-optimized paragraph-based strategy."""
        chunks = []
        document_id = self._generate_document_id(document_url)
        chunk_index = 0
        
        paragraphs = self._split_into_paragraphs(text)
        
        current_chunk = ""
        current_chunk_paragraphs = []
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
                
            # Compress paragraph if enabled
            if self.enable_chunk_compression:
                # Call the new smarter compression method
                paragraph = self._compress_text(paragraph)
            
            para_token_count = len(paragraph.split()) * 0.75
            current_token_count = len(current_chunk.split()) * 0.75
            
            if current_token_count > 0 and (
                current_token_count + para_token_count > self.max_tokens_per_chunk or
                len(current_chunk.split()) + len(paragraph.split()) > self.max_paragraph_chunk_size
            ):
                if current_chunk.strip() and len(current_chunk.split()) >= self.min_chunk_size:
                    chunk = self._create_chunk(
                        current_chunk.strip(), 
                        document_id, 
                        chunk_index, 
                        document_url
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                
                if self.chunk_overlap > 0 and current_chunk_paragraphs:
                    overlap_context = self._get_overlap_context(current_chunk_paragraphs[-1])
                    current_chunk = overlap_context + "\n\n" + paragraph if overlap_context else paragraph
                else:
                    current_chunk = paragraph
                current_chunk_paragraphs = [paragraph]
            else:
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph
                current_chunk_paragraphs.append(paragraph)
        
        if current_chunk.strip() and len(current_chunk.split()) >= self.min_chunk_size:
            chunk = self._create_chunk(
                current_chunk.strip(), 
                document_id, 
                chunk_index, 
                document_url
            )
            chunks.append(chunk)
        
        total_tokens = sum(len(chunk.text.split()) * 0.75 for chunk in chunks)
        avg_tokens_per_chunk = total_tokens / len(chunks) if chunks else 0
        
        logger.info("Created cost-optimized chunks", 
                    total_chunks=len(chunks),
                    estimated_total_tokens=int(total_tokens),
                    avg_tokens_per_chunk=int(avg_tokens_per_chunk),
                    max_tokens_per_chunk=max(len(chunk.text.split()) * 0.75 for chunk in chunks) if chunks else 0)
        
        return chunks

    def _compress_text(self, text: str) -> str:
        """
        Compresses text by replacing long, redundant phrases with shorter equivalents,
        significantly reducing token count while preserving meaning.
        """
        # A dictionary mapping long phrases (as regex) to their shorter replacements.
        # This is more powerful than just removing text.
        replacement_map = {
            # General Legal & Insurance Phrases
            r'\bwhich shall be the basis of this contract and is deemed to be incorporated herein\b': '[part of contract]',
            r'\bfollowing the Medical Advice of a duly qualified Medical Practitioner\b': 'on a doctor\'s advice',
            r'\bThe Company shall indemnify the Hospital or the Insured, Reasonable and Customary Charges incurred for Medically Necessary Treatment\b': 'Company will pay for approved medical costs',
            r'\bThe Company shall not be liable to make any payment by the Policy, in respect of any expenses incurred in connection with or in respect of\b': 'Policy excludes payment for',
            r'\bsubject to the Definitions, Terms, Exclusions, Conditions contained herein and limits\b': 'subject to policy terms and limits',
            r'\bhas applied to National Insurance Company Ltd\. \(hereinafter called the Company\)\b': 'has applied to the Company',
            r'\bsudden, unforeseen and involuntary event caused by external, visible and violent means\b': '[definition of Accident]',
            r'\bunder the supervision of a registered and qualified medical practitioner\b': 'under a qualified doctor\'s supervision',
            r'\b(shall be|are) accessible to the insurance company\'s authorized representative\b': 'accessible to the insurer',
            
            # Specific recurring clauses
            r'\bIn the event of hospitalisation/ domiciliary hospitalisation, the insured person/insured person\'s representative shall notify\b': 'For hospitalization, insured must notify',
            r'\b(for|under) any of the following circumstances\b': 'if:',
            r'\bThe services offered by a TPA shall not include\b': 'TPA services exclude:',
            r'\bThe policy shall be void and all premium paid thereon shall be forfeited to the Company\b': 'Policy will be voided',
            r'\bin the event of misrepresentation, mis description or non-disclosure of any material fact\b': 'for any non-disclosure',

            # Repeated Header/Footer info (can be removed entirely)
            r'National Insurance Co\. Ltd\.': '',
            r'Premises No\. 18-0374, Plot no\. CBD-81, New Town, Kolkata - 700156': '',
            r'National Parivar Mediclaim Plus Policy': '',
            r'UIN: NICHLIP25039V032425': '',
            r'Page \d+ of \d+': ''
        }
        
        # Apply all replacements
        for pattern, replacement in replacement_map.items():
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
            
        # Final cleanup of excessive whitespace
        text = re.sub(r'\s+', ' ', text).strip()
        
        return text

    def _get_overlap_context(self, paragraph: str) -> str:
        """Get smart overlap context - last sentence instead of full paragraph."""
        sentences = paragraph.split('.')
        if len(sentences) > 1:
            return sentences[-2].strip() + '.' if sentences[-2].strip() else ''
        return paragraph[:100] + '...' if len(paragraph) > 100 else paragraph

    def _split_into_paragraphs(self, text: str) -> List[str]:
        """Split text into meaningful paragraphs with enhanced logic."""
        paragraphs = text.split('\n\n')
        
        refined_paragraphs = []
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
                
            if len(para.split()) > self.chunk_size * 2:
                sentences = self._split_into_sentences(para)
                current_para = ""
                
                for sentence in sentences:
                    if not current_para:
                        current_para = sentence
                    elif len((current_para + " " + sentence).split()) <= self.chunk_size * 2:
                        current_para += " " + sentence
                    else:
                        refined_paragraphs.append(current_para)
                        current_para = sentence
                
                if current_para:
                    refined_paragraphs.append(current_para)
            else:
                refined_paragraphs.append(para)
        
        return refined_paragraphs
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences with improved logic."""
        sentence_pattern = r'(?<=[.!?])\s+(?=[A-Z])|(?<=[.!?])\s*\n+\s*(?=[A-Z])'
        sentences = re.split(sentence_pattern, text)
        
        cleaned_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if sentence and len(sentence.split()) >= 3:
                cleaned_sentences.append(sentence)
        
        return cleaned_sentences
    
    def _create_chunk(self, text: str, document_id: str, chunk_index: int, document_url: str) -> EmbeddingChunk:
        """Create a single embedding chunk with metadata."""
        return EmbeddingChunk(
            chunk_id=f"{document_id}_chunk_{chunk_index}",
            document_id=document_id,
            text=text,
            chunk_index=chunk_index,
            metadata={
                "page_number": self._extract_page_number(text),
                "section": self._extract_section(text),
                "clause_type": self._classify_clause_type(text),
                "document_url": document_url
            }
        )
    
    def _generate_document_id(self, url: str) -> str:
        """Generate a unique document ID from URL."""
        import hashlib
        return hashlib.md5(url.encode()).hexdigest()[:12]
    
    def _extract_page_number(self, text: str) -> Optional[int]:
        """Extract page number from chunk text."""
        page_match = re.search(r'--- Page (\d+) ---', text)
        if page_match:
            return int(page_match.group(1))
        return None
    
    def _extract_section(self, text: str) -> Optional[str]:
        """Extract section information from chunk text with enhanced logic."""
        lines = text.split('\n')
        
        for line in lines[:5]:
            line = line.strip()
            if not line:
                continue
                
            section_patterns = [
                r'^(SECTION|Section)\s+[IVX\d]+[\.\:\-\s]*(.+)',
                r'^(ARTICLE|Article)\s+[IVX\d]+[\.\:\-\s]*(.+)',
                r'^(CLAUSE|Clause)\s+[IVX\d]+[\.\:\-\s]*(.+)',
                r'^(CHAPTER|Chapter)\s+[IVX\d]+[\.\:\-\s]*(.+)',
                r'^([IVX\d]+[\.\)]\s*.{5,50})',
                r'^([A-Z][A-Z\s]{10,50}):',
            ]
            
            for pattern in section_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    return match.group(0).strip()
        
        insurance_sections = [
            'coverage', 'benefits', 'exclusions', 'limitations', 'definitions',
            'waiting period', 'grace period', 'claims', 'premium', 'renewal',
            'policy terms', 'conditions', 'procedures', 'eligibility'
        ]
        
        text_lower = text.lower()
        for section in insurance_sections:
            if section in text_lower:
                for line in lines[:3]:
                    if section in line.lower():
                        return line.strip()
                        
        return None
    
    def _classify_clause_type(self, text: str) -> Optional[str]:
        """Classify the type of clause based on content with enhanced patterns."""
        text_lower = text.lower()
        
        classification_patterns = {
            'coverage_clause': ['cover', 'benefit', 'include', 'eligible', 'entitle', 'reimburse', 'payable', 'treatment covered', 'medical expenses', 'hospital benefit'],
            'exclusion_clause': ['exclude', 'not cover', 'limitation', 'restrict', 'prohibit', 'except', 'does not include', 'not eligible', 'not payable'],
            'condition_clause': ['condition', 'require', 'must', 'shall', 'obligation', 'duty', 'responsibility', 'comply', 'fulfill', 'subject to'],
            'payment_clause': ['premium', 'payment', 'cost', 'fee', 'charge', 'amount', 'installment', 'due', 'payable', 'billing'],
            'time_clause': ['waiting period', 'grace period', 'time', 'duration', 'deadline', 'within', 'before', 'after', 'days', 'months', 'years'],
            'claims_clause': ['claim', 'settlement', 'procedure', 'process', 'submit', 'documentation', 'proof', 'evidence', 'notification'],
            'definitions_clause': ['means', 'defined as', 'definition', 'interpret', 'refer to', 'shall mean', 'is defined', 'for the purpose'],
            'renewal_clause': ['renewal', 'renew', 'extend', 'continuation', 'expiry', 'terminate', 'cancellation', 'policy period']
        }
        
        category_scores = {}
        for category, keywords in classification_patterns.items():
            score = 0
            for keyword in keywords:
                if keyword in text_lower:
                    score += len(keyword.split())
            category_scores[category] = score
        
        if category_scores:
            best_category = max(category_scores, key=category_scores.get)
            if category_scores[best_category] > 0:
                return best_category
        
        return 'general_clause'